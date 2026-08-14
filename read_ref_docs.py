"""读取参考文档内容"""
import os
import sys

try:
    from docx import Document
except ImportError:
    os.system("pip install python-docx -q")
    from docx import Document

REF_DIR = r"C:\Users\1\Desktop\参考文件"

def read_docx(filename):
    path = os.path.join(REF_DIR, filename)
    if not os.path.exists(path):
        print(f"文件不存在: {path}")
        return ""
    
    doc = Document(path)
    content = []
    for para in doc.paragraphs:
        if para.text.strip():
            content.append(para.text.strip())
    
    for table in doc.tables:
        for row in table.rows:
            row_texts = [cell.text.strip() for cell in row.cells]
            content.append(" | ".join(row_texts))
    
    return "\n".join(content)

def main():
    docs_to_read = [
        "SVG制图规范_代码落地闭环版.docx",
        "SVG制图规范v1.docx",
        "LINE215缺陷清单.docx",
        "LINE216缺陷清单.docx",
        "附件：参考内容.docx",
        "1组阶段任务书.docx",
    ]
    
    for fname in docs_to_read:
        print(f"\n{'='*80}")
        print(f"📄 读取: {fname}")
        print(f"{'='*80}")
        content = read_docx(fname)
        if content:
            print(content[:3000])
            if len(content) > 3000:
                print(f"\n... [共 {len(content)} 字符，仅显示前3000]")
        else:
            print("（空文档或读取失败）")

if __name__ == "__main__":
    main()